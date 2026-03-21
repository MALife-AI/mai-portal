#!/usr/bin/env python3
"""HWPX/OWPML → DOCX 변환기

HWPX(OWPML) 파일의 XML 구조를 파싱하여 python-docx를 이용해 DOCX로 변환합니다.

지원 요소:
  - 문단 텍스트 및 정렬 (JUSTIFY, CENTER, LEFT, RIGHT)
  - 글자 서식 (폰트 크기, 굵게)
  - 줄간격
  - 표 (셀 병합 포함)
  - 수식 (OMML 네이티브 수식으로 변환)
  - 페이지 설정 (용지 크기, 여백)

사용법:
  python3 hwpx2docx.py input.hwpx [output.docx]
  python3 hwpx2docx.py input.owpml [output.docx]
"""

import sys
import os
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from backend.ingestion.converters.hwpeq2omml import hwp_eq_to_omml

NS = {
    'hp': 'http://www.hancom.co.kr/hwpml/2011/paragraph',
    'hs': 'http://www.hancom.co.kr/hwpml/2011/section',
    'hc': 'http://www.hancom.co.kr/hwpml/2011/core',
    'hh': 'http://www.hancom.co.kr/hwpml/2011/head',
    'opf': 'http://www.idpf.org/2007/opf/',
}

# HWPX 단위: 1/7200 inch, 1 inch = 914400 EMU
HWPUNIT_TO_EMU = 914400 / 7200  # = 127


class HwpxParser:
    def __init__(self, filepath):
        self.filepath = filepath
        self.zf = zipfile.ZipFile(filepath, 'r')
        self.char_props = {}  # id -> {height, bold, font_hangul_id, font_latin_id}
        self.para_props = {}  # id -> {align, line_spacing_type, line_spacing_value}
        self.fonts_hangul = {}  # id -> face name
        self.fonts_latin = {}
        self.border_fills = {}  # id -> border info
        self._parse_header()

    def _parse_header(self):
        header_xml = self.zf.read('Contents/header.xml')
        root = ET.fromstring(header_xml)

        # 폰트 파싱
        for fontface in root.findall('.//hh:fontface', NS):
            lang = fontface.get('lang', '')
            font_dict = self.fonts_hangul if lang == 'HANGUL' else self.fonts_latin if lang == 'LATIN' else {}
            for font in fontface.findall('hh:font', NS):
                fid = font.get('id')
                face = font.get('face', '')
                if fid is not None:
                    font_dict[int(fid)] = face

        # 글자 속성 파싱
        for cp in root.findall('.//hh:charPr', NS):
            cpid = int(cp.get('id', '-1'))
            height = int(cp.get('height', '1000'))
            bold = cp.find('hh:bold', NS) is not None

            font_hangul_id = None
            font_latin_id = None
            fref = cp.find('hh:fontRef', NS)
            if fref is not None:
                h = fref.get('hangul')
                l = fref.get('latin')
                if h is not None:
                    font_hangul_id = int(h)
                if l is not None:
                    font_latin_id = int(l)

            self.char_props[cpid] = {
                'height': height,
                'bold': bold,
                'font_hangul_id': font_hangul_id,
                'font_latin_id': font_latin_id,
            }

        # 문단 속성 파싱
        for pp in root.findall('.//hh:paraPr', NS):
            ppid = int(pp.get('id', '-1'))
            align_el = pp.find('hh:align', NS)
            align = align_el.get('horizontal', 'LEFT') if align_el is not None else 'LEFT'

            ls_el = pp.find('hh:lineSpacing', NS)
            ls_type = ls_el.get('type', 'PERCENT') if ls_el is not None else 'PERCENT'
            ls_value = ls_el.get('value', '150') if ls_el is not None else '150'

            margin_el = pp.find('hh:margin', NS)
            indent = 0
            left = 0
            right = 0
            if margin_el is not None:
                indent = int(margin_el.get('indent', '0'))
                left = int(margin_el.get('left', '0'))
                right = int(margin_el.get('right', '0'))

            self.para_props[ppid] = {
                'align': align,
                'line_spacing_type': ls_type,
                'line_spacing_value': int(ls_value),
                'indent': indent,
                'left': left,
                'right': right,
            }

        # 테두리/채우기 파싱
        for bf in root.findall('.//hh:borderFill', NS):
            bfid = int(bf.get('id', '-1'))
            borders = {}
            for side in ['leftBorder', 'rightBorder', 'topBorder', 'bottomBorder']:
                el = bf.find(f'hh:{side}', NS)
                if el is not None:
                    borders[side] = {
                        'type': el.get('type', 'NONE'),
                        'width': el.get('width', '0.1 mm'),
                        'color': el.get('color', '#000000'),
                    }
            self.border_fills[bfid] = borders

    def get_sections(self):
        """content.hpf에서 섹션 파일 목록을 가져옴"""
        content_hpf = self.zf.read('Contents/content.hpf')
        root = ET.fromstring(content_hpf)
        sections = []
        for item in root.findall('.//opf:itemref', NS):
            idref = item.get('idref', '')
            if idref.startswith('section'):
                sections.append(f'Contents/{idref}.xml')
        return sections

    def parse_section(self, section_path):
        """섹션 XML을 파싱하여 요소 리스트를 반환"""
        section_xml = self.zf.read(section_path)
        root = ET.fromstring(section_xml)
        elements = []
        self._parse_children(root, elements)
        return elements, root

    def _parse_children(self, parent, elements):
        for child in parent:
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                self._parse_paragraph(child, elements)
            elif tag == 'tbl':
                self._parse_table(child, elements)

    def _parse_paragraph(self, p_el, elements):
        para_pr_id = int(p_el.get('paraPrIDRef', '0'))

        # secPr 확인 (페이지 설정)
        sec_pr = p_el.find('.//hp:secPr', NS)
        if sec_pr is not None:
            page_pr = sec_pr.find('hp:pagePr', NS)
            margin_pr = sec_pr.find('.//hp:margin', NS)
            if page_pr is not None:
                elements.append({
                    'type': 'page_setup',
                    'width': int(page_pr.get('width', '59532')),
                    'height': int(page_pr.get('height', '84200')),
                    'landscape': page_pr.get('landscape', 'WIDELY'),
                    'margin': {
                        'top': int(margin_pr.get('top', '2835')) if margin_pr is not None else 2835,
                        'bottom': int(margin_pr.get('bottom', '2835')) if margin_pr is not None else 2835,
                        'left': int(margin_pr.get('left', '5386')) if margin_pr is not None else 5386,
                        'right': int(margin_pr.get('right', '5386')) if margin_pr is not None else 5386,
                        'header': int(margin_pr.get('header', '2835')) if margin_pr is not None else 2835,
                        'footer': int(margin_pr.get('footer', '2835')) if margin_pr is not None else 2835,
                    },
                })

        # 표가 포함된 문단은 표로 처리
        tbl = p_el.find('.//hp:tbl', NS)
        if tbl is not None:
            self._parse_table(tbl, elements)
            return

        runs = []
        for run_el in p_el.findall('hp:run', NS):
            char_pr_id = int(run_el.get('charPrIDRef', '0'))

            # 수식 체크
            eq = run_el.find('hp:equation', NS)
            if eq is not None:
                script_el = eq.find('hp:script', NS)
                eq_script = ''
                if script_el is not None and script_el.text:
                    eq_script = script_el.text.strip()
                if not eq_script:
                    eq_script = eq.get('script', '')
                if eq_script:
                    runs.append({
                        'type': 'equation',
                        'script': eq_script,
                        'char_pr_id': char_pr_id,
                    })
                continue

            # ctrl 안에 표가 있는 경우
            ctrl = run_el.find('hp:ctrl', NS)
            if ctrl is not None:
                inner_tbl = ctrl.find('.//hp:tbl', NS)
                if inner_tbl is not None:
                    # 현재까지의 runs가 있으면 문단으로 먼저 추가
                    if runs:
                        elements.append({
                            'type': 'paragraph',
                            'para_pr_id': para_pr_id,
                            'runs': runs,
                        })
                        runs = []
                    self._parse_table(inner_tbl, elements)
                    continue

            # 텍스트 수집
            for t_el in run_el.findall('hp:t', NS):
                text = t_el.text or ''
                if text:
                    runs.append({
                        'text': text,
                        'char_pr_id': char_pr_id,
                    })

        if runs:
            elements.append({
                'type': 'paragraph',
                'para_pr_id': para_pr_id,
                'runs': runs,
            })
        elif tbl is None:
            # 빈 문단
            elements.append({
                'type': 'paragraph',
                'para_pr_id': para_pr_id,
                'runs': [],
            })

    def _parse_table(self, tbl_el, elements):
        table_data = {
            'type': 'table',
            'rows': [],
            'col_cnt': int(tbl_el.get('colCnt', '1')),
            'border_fill_id': int(tbl_el.get('borderFillIDRef', '1')),
        }

        for tr_el in tbl_el.findall('hp:tr', NS):
            row = []
            for tc_el in tr_el.findall('hp:tc', NS):
                cell_addr = tc_el.find('hp:cellAddr', NS)
                cell_span = tc_el.find('hp:cellSpan', NS)
                cell_sz = tc_el.find('hp:cellSz', NS)

                col_addr = int(cell_addr.get('colAddr', '0')) if cell_addr is not None else 0
                row_addr = int(cell_addr.get('rowAddr', '0')) if cell_addr is not None else 0
                col_span = int(cell_span.get('colSpan', '1')) if cell_span is not None else 1
                row_span = int(cell_span.get('rowSpan', '1')) if cell_span is not None else 1
                width = int(cell_sz.get('width', '0')) if cell_sz is not None else 0
                height = int(cell_sz.get('height', '0')) if cell_sz is not None else 0

                # 셀 내 문단들
                cell_paras = []
                for cp in tc_el.findall('.//hp:p', NS):
                    cell_para_pr_id = int(cp.get('paraPrIDRef', '0'))
                    cell_runs = []
                    for run_el in cp.findall('hp:run', NS):
                        char_pr_id = int(run_el.get('charPrIDRef', '0'))
                        # 수식 체크
                        eq = run_el.find('hp:equation', NS)
                        if eq is not None:
                            script_el = eq.find('hp:script', NS)
                            eq_script = ''
                            if script_el is not None and script_el.text:
                                eq_script = script_el.text.strip()
                            if eq_script:
                                cell_runs.append({
                                    'type': 'equation',
                                    'script': eq_script,
                                    'char_pr_id': char_pr_id,
                                })
                            continue
                        for t_el in run_el.findall('hp:t', NS):
                            text = t_el.text or ''
                            if text:
                                cell_runs.append({
                                    'text': text,
                                    'char_pr_id': char_pr_id,
                                })
                    cell_paras.append({
                        'para_pr_id': cell_para_pr_id,
                        'runs': cell_runs,
                    })

                border_fill_id = int(tc_el.get('borderFillIDRef', '1')) if tc_el.get('borderFillIDRef') else 1

                row.append({
                    'col_addr': col_addr,
                    'row_addr': row_addr,
                    'col_span': col_span,
                    'row_span': row_span,
                    'width': width,
                    'height': height,
                    'paragraphs': cell_paras,
                    'border_fill_id': border_fill_id,
                })
            table_data['rows'].append(row)

        elements.append(table_data)

    def close(self):
        self.zf.close()


class DocxWriter:
    ALIGN_MAP = {
        'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
        'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
        'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
        'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    def __init__(self, parser: HwpxParser):
        self.parser = parser
        self.doc = Document()

    def _hwpunit_to_emu(self, val):
        return int(val * HWPUNIT_TO_EMU)

    def _apply_page_setup(self, setup_info):
        section = self.doc.sections[-1]
        section.page_width = self._hwpunit_to_emu(setup_info['width'])
        section.page_height = self._hwpunit_to_emu(setup_info['height'])

        m = setup_info['margin']
        section.top_margin = self._hwpunit_to_emu(m['top'])
        section.bottom_margin = self._hwpunit_to_emu(m['bottom'])
        section.left_margin = self._hwpunit_to_emu(m['left'])
        section.right_margin = self._hwpunit_to_emu(m['right'])
        section.header_distance = self._hwpunit_to_emu(m['header'])
        section.footer_distance = self._hwpunit_to_emu(m['footer'])

    def _apply_para_format(self, paragraph, para_pr_id):
        pp = self.parser.para_props.get(para_pr_id, {})

        # 정렬
        align = pp.get('align', 'LEFT')
        paragraph.alignment = self.ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)

        # 줄간격
        pf = paragraph.paragraph_format
        ls_type = pp.get('line_spacing_type', 'PERCENT')
        ls_value = pp.get('line_spacing_value', 150)

        if ls_type == 'PERCENT':
            from docx.shared import Pt as _Pt
            pf.line_spacing = ls_value / 100.0

        # 들여쓰기
        indent = pp.get('indent', 0)
        left = pp.get('left', 0)
        right = pp.get('right', 0)
        if indent:
            pf.first_line_indent = self._hwpunit_to_emu(indent)
        if left:
            pf.left_indent = self._hwpunit_to_emu(left)
        if right:
            pf.right_indent = self._hwpunit_to_emu(right)

    def _apply_run_format(self, run, char_pr_id):
        cp = self.parser.char_props.get(char_pr_id, {})

        # 폰트 크기 (height 단위: 1/100 pt)
        height = cp.get('height', 1000)
        run.font.size = Pt(height / 100.0)

        # 굵게
        if cp.get('bold'):
            run.font.bold = True

        # 폰트 이름
        font_hangul_id = cp.get('font_hangul_id')
        font_latin_id = cp.get('font_latin_id')

        font_name = None
        if font_hangul_id is not None and font_hangul_id in self.parser.fonts_hangul:
            font_name = self.parser.fonts_hangul[font_hangul_id]
        elif font_latin_id is not None and font_latin_id in self.parser.fonts_latin:
            font_name = self.parser.fonts_latin[font_latin_id]

        if font_name:
            run.font.name = font_name
            # 한글 폰트 설정
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = ET.SubElement(rpr, qn('w:rFonts'))
            rfonts.set(qn('w:eastAsia'), font_name)

    def _add_paragraph(self, para_info):
        p = self.doc.add_paragraph()
        self._apply_para_format(p, para_info['para_pr_id'])

        for run_info in para_info.get('runs', []):
            if run_info.get('type') == 'equation':
                self._insert_equation(p, run_info['script'])
            else:
                run = p.add_run(run_info['text'])
                self._apply_run_format(run, run_info['char_pr_id'])

        return p

    def _insert_equation(self, paragraph, script):
        """수식 스크립트를 OMML로 변환하여 문단에 삽입"""
        from backend.ingestion.converters.hwpeq2omml import omml_to_string
        from lxml import etree as lxml_etree
        omath = hwp_eq_to_omml(script)
        # stdlib ET → lxml 변환
        xml_str = omml_to_string(omath)
        lxml_omath = lxml_etree.fromstring(xml_str)
        paragraph._element.append(lxml_omath)

    def _add_table(self, table_info):
        rows_data = table_info['rows']
        if not rows_data:
            return

        # 행/열 수 계산
        num_rows = len(rows_data)
        num_cols = table_info.get('col_cnt', 1)

        # 실제 열 수를 셀 주소 + span에서 재계산
        max_col = 0
        for row in rows_data:
            for cell in row:
                end_col = cell['col_addr'] + cell['col_span']
                if end_col > max_col:
                    max_col = end_col
        if max_col > num_cols:
            num_cols = max_col

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 기본 테이블 스타일
        table.style = 'Table Grid'

        # 셀 병합 및 내용 채우기
        merged_cells = set()

        for ri, row_data in enumerate(rows_data):
            for cell_info in row_data:
                col_addr = cell_info['col_addr']
                row_addr = cell_info['row_addr']
                col_span = cell_info['col_span']
                row_span = cell_info['row_span']

                if (ri, col_addr) in merged_cells:
                    continue

                try:
                    cell = table.cell(ri, col_addr)
                except IndexError:
                    continue

                # 셀 병합
                if col_span > 1 or row_span > 1:
                    end_row = min(ri + row_span - 1, num_rows - 1)
                    end_col = min(col_addr + col_span - 1, num_cols - 1)
                    try:
                        merge_cell = table.cell(end_row, end_col)
                        cell.merge(merge_cell)
                        for mr in range(ri, end_row + 1):
                            for mc in range(col_addr, end_col + 1):
                                if (mr, mc) != (ri, col_addr):
                                    merged_cells.add((mr, mc))
                    except IndexError:
                        pass

                # 셀 내용
                cell_paras = cell_info.get('paragraphs', [])
                if cell_paras:
                    # 첫 번째 문단은 기존 셀 문단 사용
                    for pi, cp in enumerate(cell_paras):
                        if pi == 0:
                            p = cell.paragraphs[0]
                        else:
                            p = cell.add_paragraph()

                        self._apply_para_format(p, cp['para_pr_id'])

                        for run_info in cp.get('runs', []):
                            if run_info.get('type') == 'equation':
                                self._insert_equation(p, run_info['script'])
                            else:
                                run = p.add_run(run_info['text'])
                                self._apply_run_format(run, run_info['char_pr_id'])

                # 셀 너비 설정
                if cell_info['width'] > 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = ET.SubElement(tcPr, qn('w:tcW'))
                    # HWPX width를 twips로 변환 (1 hwpunit = 127 EMU, 1 twip = 635 EMU)
                    width_twips = int(cell_info['width'] * 127 / 635)
                    tcW.set(qn('w:w'), str(width_twips))
                    tcW.set(qn('w:type'), 'dxa')

    def convert(self, output_path):
        sections = self.parser.get_sections()

        for section_path in sections:
            elements, root = self.parser.parse_section(section_path)

            for el in elements:
                if el['type'] == 'page_setup':
                    self._apply_page_setup(el)
                elif el['type'] == 'paragraph':
                    self._add_paragraph(el)
                elif el['type'] == 'table':
                    self._add_table(el)

        self.doc.save(output_path)
        print(f'변환 완료: {output_path}')


def main():
    if len(sys.argv) < 2:
        print(f'사용법: python3 {sys.argv[0]} <input.hwpx|input.owpml> [output.docx]')
        sys.exit(1)

    input_path = sys.argv[1]
    if not os.path.exists(input_path):
        print(f'파일을 찾을 수 없습니다: {input_path}')
        sys.exit(1)

    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        base = os.path.splitext(input_path)[0]
        output_path = base + '.docx'

    parser = HwpxParser(input_path)
    writer = DocxWriter(parser)
    writer.convert(output_path)
    parser.close()


if __name__ == '__main__':
    main()
